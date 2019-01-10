# Import statements
import re

import nltk
from docx import Document
from nltk import word_tokenize
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from IndeedAPI import *


class DataProcessor:
    # function to get paragraphs / sections from word document
    def gen_paragraph(self, document):
        # holds the document paragraphs
        resume = [line.text for line in document.paragraphs]

        # True/False array
        isEmptyArr = []

        # Array that will hold the paragraphs
        paragraphs = []

        section = []
        # array that will populate the true false array
        for line in resume:
            if len(line) == 0 or bool(re.match('^\s+$', line)) or line in '()~><?'';":\/|{},[]@6&*-_.':
                isEmptyArr.append(True)
            else:
                isEmptyArr.append(False)

        for i in range(len(resume)):
            # print("%s " % resume[i] + "||| % s" % isEmptyArr[i])
            if isEmptyArr[i] is not True:
                section += resume[i]
            else:
                paragraphs.append(section)
                section = ""
        paragraphs.append(section)
        return paragraphs

    # input: job title as a string, location as int, page limit as int
    # Returns a dictionary of {job_id: {title: "string", description: "string", keywords: ["String"]}}
    def get_jobs(self, title, location, page_limit):
        indeedApi = IndeedAPi()
        job_urls = indeedApi.retrieve_urls(title, location, page_limit)
        # time.sleep(15)
        # this list holds jobs that may have no content.
        unfiltered_jobs = [indeedApi.getJobMarkup(url) for url in job_urls]

        # filtered jobs are jobs in list that have content Note: this holds the markup (html) of the job
        filtered_jobs = indeedApi.filterJobs(unfiltered_jobs)

        # takes the filtered jobs and creates a dictionary that holds all the jobs
        # key = {job_id: {title: "string", description: "string", keywords: ["String"]}}
        json_data = indeedApi.generateJson(filtered_jobs)

        # generate the data using paragraph separation
        json_data_wparagraphs = indeedApi.generateJsonParagraphs(filtered_jobs)
        return json_data, json_data_wparagraphs


    # takes in resume and bool to indicate whether to return list of paragraphs or the entire resume as a string
    def process_resume(self, resume, split_paragraphs):
        file = open(resume, 'rb')
        d = Document(file)
        fullText = []
        resumeText = ""

        for p in d.paragraphs:
            fullText.append(p.text)
        if split_paragraphs is True:
            paragraphs = self.gen_paragraph(d)
            return paragraphs
        else:
            # Iterate through the array removing indices that have length of 0
            for text in fullText:
                try:
                    if len(text) == 0:
                        fullText.remove(text)
                except IndexError:
                    print("Could not pop indice...")
            # iterate through the fullText array and strip \n and \t
            for i in range(len(fullText)):
                fullText[i] = re.sub("\s+", " ",
                                     fullText[i])  # if there is text remove the extra space and the and the tabs
                resumeText = resumeText + fullText[
                    i] + " "  # concatinate part of the resume text to the job description
            return resumeText

    # pass in a string or list of paragraphs
    # Strips the stopwords for the resume and removes any punctuation that may be found at the last chr of the str
    # remove parts of speech
    def strip_resume_stopwords_punctuation_pos(self, resume):
        punctuation = '~><?'';:\/|{}[]@6&*-_,.'
        POSToKeep = ["JJ", "JJR", "JJS", "NN", "NNS", "NNP", "NNPS", "VB", "VBD", "VBG", "VBN", "VBP",
                     "VBZ"]  # These are the nouns, verbs, and adjectives we want to keep

        # if string is passed, then return a list of words from the resume with no stopwords
        if isinstance(resume, str):
            resumeNoStopWords = [word for word in resume.split() if word not in stopwords.words('english')]

            # snippet that strips punctuation from the words in the resume.
            for i in range(len(resumeNoStopWords)):
                if resumeNoStopWords[i][-1] in punctuation:
                    resumeNoStopWords[i] = resumeNoStopWords[i].replace(resumeNoStopWords[i][-1], "")
            # iterate through the words in the resume and remove the terms that are less than 2 and contain any special characters
            for i in range(len(resumeNoStopWords)):
                try:
                    if len(resumeNoStopWords[i]) < 2 and resumeNoStopWords[i][0] in '()~><?'';":\/|{},[]@6&*-_':
                        resumeNoStopWords.remove(resumeNoStopWords[i])
                except IndexError:
                    print("Could not pop indice...")
            # retrieve parts of speech from each term in the resume, keeping only nouns verbs and adjectives
            resumeNoStopWords = " ".join(w for w in resumeNoStopWords if len(w) > 0)

            resumePOSTags = nltk.pos_tag(resumeNoStopWords.split())
            for w in resumePOSTags:
                if w[1] not in POSToKeep:
                    resumePOSTags.remove(w)
            resumeNoStopWordsUpdated = [w[0] for w in resumePOSTags]
            return resumeNoStopWordsUpdated


        # if list is passed in, then return a 2d list, which are the paragraphs consisting of stopwords
        elif isinstance(resume, list):
            resumeNoStopWords = []
            for p in resume:
                if len(p) < 1:
                    continue
                else:
                    # Strip the stopwords
                    # add each word to the paragraph and add the paragraph to the list
                    paragraph = [word for word in p.split() if word not in stopwords.words('english')]
                    resumeNoStopWords.append(paragraph)

            # Remove puncuation
            for i in range(len(resumeNoStopWords)):
                for j in range(len(resumeNoStopWords[i])):
                    if resumeNoStopWords[i][j][-1] in punctuation:
                        resumeNoStopWords[i][j] = resumeNoStopWords[i][j].replace(resumeNoStopWords[i][j][-1], "")

            # iterate through the words in the resume and remove the terms that are less than 2 and contain any special characters
            for i in range(len(resumeNoStopWords)):
                for j in range(len(resumeNoStopWords[i])):
                    try:
                        if len(resumeNoStopWords[i][j]) < 2 and resumeNoStopWords[i][j][
                            0] in '()~><?'';":\/|{},[]@6&*-_':
                            resumeNoStopWords[i].remove(resumeNoStopWords[i][j])
                    except IndexError:
                        print("Could not pop indice...")

            resumePOSTags = [nltk.pos_tag(word_tokenize(" ".join(w for w in section))) for section in resumeNoStopWords]
            for i, line in enumerate(resumePOSTags):
                for w in line:
                    if w[1] not in POSToKeep:
                        resumePOSTags[i].remove(w)

            resumeNoStopWordsUpdated = [[w[0] for w in line] for line in resumePOSTags]
            return resumeNoStopWordsUpdated



        else:
            return Exception("list or str not passed in")

    # Pass in jobs dictionary, strip stopwords, remove certain parts of speech, and strip punctuation
    def process_jobs(self, jobs):
        containsNumsSpecialChars = r'^[!@#$%^&*(),.?":{}|<>0-9]*$'
        POSToKeep = ["JJ", "JJR", "JJS", "NN", "NNS", "NNP", "NNPS", "VB", "VBD", "VBG", "VBN", "VBP",
                     "VBZ"]  # These are the nouns, verbs, and adjectives we want to keep
        punctuation = '~><?'';:\/|{}[]@6&*-_,.'
        jobsNoStopWords = {}  # job_id: {title: "string", descriptionNoStopwords: []
        # job_id = 0
        for i in range(len(jobs)):
            # Create an empty dictionary
            jobsNoStopWords[i] = {}

            # Store title from json_data
            jobsNoStopWords[i]["title"] = jobs[i]["title"]

            # Store the words with the stopwords stripped out represented as a list
            jobsNoStopWords[i]["description"] = [w for w in jobs[i]["description"].split() if
                                                 w not in stopwords.words('english')]
            # Filter out terms containing both numbers and special characters
            jobsNoStopWords[i]["description"] = [w for w in jobs[i]["description"].split() if
                                                 bool(re.match(containsNumsSpecialChars, w)) is False]

        jobsNoStopWordsUpdated = {}

        # Remove parts of speech
        for i in range(len(jobsNoStopWords)):

            jobsNoStopWordsUpdated[i] = {}
            jobsNoStopWordsUpdated[i]["title"] = jobsNoStopWords[i]["title"]

            # store the job part of speech
            jobPOS = nltk.pos_tag(jobsNoStopWords[i]["description"])
            for w in jobPOS:
                if w[1] not in POSToKeep:
                    print("found %s" % w[1])
                    jobPOS.remove(w)
            jobsNoStopWordsUpdated[i]["description"] = [w[0] for w in jobPOS]

        # Remove special punctuation from the last indice in the jobsNoStopWordsUpdated dictionary description key
        for i in range(len(jobsNoStopWordsUpdated)):
            for j in range(len(jobsNoStopWordsUpdated[i]["description"])):
                if jobsNoStopWordsUpdated[i]["description"][j][-1] in punctuation:
                    jobsNoStopWordsUpdated[i]["description"][j] = jobsNoStopWordsUpdated[i]["description"][j].replace(
                        jobsNoStopWordsUpdated[i]["description"][j][-1], "")
        return jobsNoStopWordsUpdated

    # Method to process the paragraph separated jobs
    def process_jobs_paragraphs(self, jobs):
        containsNumsSpecialChars = r'^[!@#$%^&*(),.?":{}|<>0-9]*$'
        POSToKeep = ["JJ", "JJR", "JJS", "NN", "NNS", "NNP", "NNPS", "VB", "VBD", "VBG", "VBN", "VBP",
                     "VBZ"]  # These are the nouns, verbs, and adjectives we want to keep
        punctuation = list('~><?'';:\/|{}[]@&*-_,.)(')
        # Process parts of speech
        for i in range(len(jobs)):
            if len(jobs[i]["description"]) < 1:
                continue
            for j in range(len(jobs[i]["description"])):
                paragraphPOS = nltk.pos_tag(jobs[i]["description"][j])
                for word in paragraphPOS:
                    if word[1] not in POSToKeep:
                        print("found %s" % word[1])
                        paragraphPOS.remove(word)
            jobs[i]["description"][j] = [word[0] for word in paragraphPOS if
                                         len(word[0]) > 1 and word[0] not in punctuation]

        # Remove stopwords
        for i in range(len(jobs)):
            for j in range(len(jobs[i]["description"])):
                jobs[i]["description"][j] = [word for word in jobs[i]["description"][j] if
                                             word not in stopwords.words('english')]
                # Remove numbers and special characters only
                jobs[i]["description"][j] = [word for word in jobs[i]["description"][j] if
                                             bool(re.match(containsNumsSpecialChars, word)) is False]

        return jobs
    # Function to determine the tf_idf
    # Training set = job descriptions
    # Test set = resume(s) or set of resumes
    # Returns tuple consisting of x,y and vocab
    def tf_idf(self, training_set, test_set):
        tf_idf_vectorizer = TfidfVectorizer(use_idf=False, sublinear_tf=False, stop_words=stopwords.words('english'))
        corpus = []
        for i in range(len(training_set)):
            corpus.append(training_set[i]["title"] + " " + " ".join(word for word in training_set[i]["description"]))

        # create bag of words and transform corpus into a document term frequency matrix
        x = tf_idf_vectorizer.fit_transform(corpus)

        # transform the resume into a term frequency matrix
        resumeCorpus = " ".join(word for word in test_set)
        y = tf_idf_vectorizer.transform([resumeCorpus])

        # This portion of code is untested

        # get the vocabulary
        features = tf_idf_vectorizer.get_feature_names()

        # return tuple
        return (x, y, features)

    # Tf-idf helper function for paragraphs in the resume to paragraphs in the job descriptions
    def tf_idf2(self, jobs, resume):
        # scores are sorted as {job#: {resume paragraph#: score}}
        total_scores = {}
        for job_num in jobs.items():
            singlejob = job_num[1]
            corpus = []

            # add paragraphs of the particular job to the corpus
            for paragraph in singlejob["description"]:
                paragraphstr = " ".join(word for word in paragraph)
                corpus.append(paragraphstr)

            # counter for resume paragraph number starting at 1
            resume_paragraph_num = 0

            # Create a nested dictionary where the key is the paragraph
            total_scores[job_num[0]] = {}
            # compute the score for each paragraph in the resume
            print(corpus)
            for paragraph in resume:
                corpusParagraph = " ".join(word for word in paragraph)
                tf_idf_vectorizer = TfidfVectorizer(use_idf=False, sublinear_tf=False,
                                                    stop_words=stopwords.words('english'))
                # x is the job set
                # y is the resume set
                x = tf_idf_vectorizer.fit_transform(corpus)
                y = tf_idf_vectorizer.transform([corpusParagraph])
                total_scores[job_num[0]][resume_paragraph_num] = (x, y)
                resume_paragraph_num += 1

        return total_scores


    # Function to get the cosine similarity
    # Takes in two document term frequency matrixes returned from the tf-idf function
    def get_cosine_similarity(self, x, y):
        return cosine_similarity(x, y).flatten()

    # Function that gets the bigrams for the document.
    def get_bigrams(self, dataset, occurences):
        bigrams = list(nltk.bigrams(dataset))
        for bi in bigrams:
            bigramstr = bi[0] + " " + bi[1]
            if bigrams.count(bi) >= occurences:
                if bigramstr not in dataset:
                    dataset.append(bigramstr)
        return dataset

    #This gets the bigrams for the entire collection of documents
    def get_all_bigrams(self,dataset,occurences):
        # list to hold all the bigrams from the entire corpus.
        all_bigrams = []
        for i in range(len(dataset)):
            bigrams = list(nltk.bigrams(dataset[i]["description"]))

            for bi in bigrams:
                bigramstr = bi[0] + " " + bi[1]
                all_bigrams.append(bigramstr)

        updated_bigrams = [bi for bi in all_bigrams if all_bigrams.count(bi) >= occurences]
        for i in range(len(dataset)):
            bigrams = list(nltk.bigrams(dataset[i]["description"]))
            for bi in bigrams:
                bigramstr = bi[0] + " " + bi[1]
                if bigramstr in updated_bigrams:
                    dataset[i]["description"].append(bigramstr)
        return dataset

    # Generates bigrams for jobs that are separated into paragraphs
    # Note: this takes in the data set where the job description is passed as a list of strings
    def get_all_bigrams_paragraphs(self, dataset, occurrences):
        all_bigrams = []

        # Iterate through the dictionary
        for i in range(len(dataset)):
            # ite1ate through the paragraphs within the jd
            jd = dataset[i]["description"]
            for j in range(len(jd)):
                bigrams = list(nltk.bigrams(word_tokenize(jd[j])))
                for bi in bigrams:
                    bigramstr = bi[0] + " " + bi[1]
                    all_bigrams.append(bigramstr)
        updated_bigrams = [bi for bi in all_bigrams if all_bigrams.count(bi) >= occurrences]

        for i in range(len(dataset)):
            for j in range(len(dataset[i]["description"])):
                dataset[i]["description"][j] = word_tokenize(dataset[i]["description"][j])

        print("adding bigrams\n")


        for i in range(len(dataset)):
            jd = dataset[i]["description"]
            for j in range(len(jd)):
                bigrams = list(nltk.bigrams(jd[j]))
                for bi in bigrams:
                    bigramstr = bi[0] + " " + bi[1]
                    if bigramstr in updated_bigrams:
                        print("appending %s" % bigramstr)
                        dataset[i]["description"][j].append(bigramstr)

            # Returns the dataset where the description is turned into a list of lists of strings
            return dataset

    # helper method to get skills from a corpus
    # input: corpus such as a resume or job description
    # Implementation not working... fix it
    # Best to pass in a preprocessed corpus
    # Consider passing in a tuple instead of a list
    def filter_pos(self, corpus, POS_to_keep):
        # List of parts of speech to keep
        # We are only wanting to keep the nouns & adjectives
        # POS_to_keep = ["NN", "NNS", "NNP", "NNPS", "JJ", "JJR", "JJS"]
        possible_skills = []
        # Check if what is being passed is a list of strings
        if isinstance(corpus, list):
            for section in corpus:
                pos_tokenized_section = nltk.pos_tag(section.split())
                for word in pos_tokenized_section:
                    if word[1] in POS_to_keep:
                        possible_skills.append(word[0])
            return possible_skills
        # Check if a single string is being passed in
        elif isinstance(corpus, str):
            possible_skills = [word[0] for word in nltk.pos_tag(corpus.split()) if word[1] in POS_to_keep]
            return possible_skills
        else:
            Exception("List or string must be passed but other type found instead.")
