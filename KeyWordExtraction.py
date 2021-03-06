# Test file that extracts

import re

import nltk
from docx import Document
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from IndeedAPI import *


#data needs to be exported to json.


def gen_paragraph(document):
    # holds the document paragraphs
    resume = [line.text for line in document.paragraphs]

    # True/False array
    isEmptyArr = []

    # Array that will hold the paragraphs
    paragraphs = []

    section = []
    # array that will populate the true false array
    for line in resume:
        if len(line) == 0 or bool(re.match('^\s+$', line)) or line in '()~><?'';":\/|{},[]@6&*-_.':
            isEmptyArr.append(True)
        else:
            isEmptyArr.append(False)

    for i in range(len(resume)):
        # print("%s " % resume[i] + "||| % s" % isEmptyArr[i])
        if isEmptyArr[i] is not True:
            section += resume[i]
        else:
            paragraphs.append(section)
            section = ""
    paragraphs.append(section)
    return paragraphs

#Initialize the necessary objects
indeedApi = IndeedAPi()

# Section to retrieve the urls and data from the jobs
job_urls = indeedApi.retrieve_urls("Java developer",11590,10)

#this list holds jobs that may have no content.
unfiltered_jobs = [indeedApi.getJobMarkup(url) for url in job_urls]

# filtered jobs are jobs in list that have content Note: this holds the markup (html) of the job
filtered_jobs = indeedApi.filterJobs(unfiltered_jobs)

#takes the filtered jobs and creates a dictionary that holds all the jobs
# key = {job_id: {title: "string", description: "string", keywords: ["String"]}}
json_data = indeedApi.generateJson(filtered_jobs)

# extract text from documents
file = open("TalWeissResume.docx", 'rb')
d = Document(file)
fullText = []

for p in d.paragraphs:
    fullText.append(p.text)

resumetext = ""  # empty string which should hold the text of the resume

paragraphs = gen_paragraph(d)
# Iterate through the array removing indices that have length of 0
# for i in range(len(fullText)):
#     try:
#         if len(fullText[i]) == 0:
#             fullText.pop(i)
#     except IndexError:
#         print("Could not pop indice...")
# iterate through the fullText array and strip \n and \t
# for j in range(len(paragraphs)):
for i in range(len(paragraphs)):
    if len(paragraphs[i]) < 1:
        continue
    else:
        paragraphs[i] = re.sub("\s+", " ",
                               paragraphs[i])  # if there is text remove the extra space and the and the tabs


# strip the stopwords from the resume
# resumeNoStopWords = [[word for word in resumetext.split() if word not in stopwords.words('english')]]
resumeNoStopWords = []
for p in paragraphs:
    if len(p) < 1:
        continue
    else:
        paragraph = [word for word in p.split() if word not in stopwords.words('english')]
        resumeNoStopWords.append(paragraph)




# iterate through the words in the resume and remove the terms that are less than 2 and contain any special characters
for i in range(len(resumeNoStopWords)):
    for j in range(len(resumeNoStopWords[i])):
        try:
            if len(resumeNoStopWords[i][j]) < 2 and resumeNoStopWords[i][j][0] in '()~><?'';":\/|{},[]@6&*-_':
                resumeNoStopWords[i].pop(j)
        except IndexError:
            print("Could not pop indice...")

# iterate through json_data job title and description and strip the stopwords from the description
jobsNoStopWords = {}  # job_id: {title: "string", descriptionNoStopwords: []
# job_id = 0
for i in range(len(json_data)):
    # Create an empty dictionary
    jobsNoStopWords[i] = {}

    # Store title from json_data
    jobsNoStopWords[i]["title"] = json_data[i]["title"]

    # Store the words with the stopwords stripped out represented as a list
    jobsNoStopWords[i]["description"] = [w for w in json_data[i]["description"].split() if
                                         w not in stopwords.words('english')]

# retrieve parts of speech from each term in the resume, keeping only nouns verbs and adjectives
# resumePOSTags = nltk.pos_tag(resumeNoStopWords)
resumePOSTags = [nltk.pos_tag(section) for section in resumeNoStopWords]
POSToKeep = ["JJ", "JJR", "JJS", "NN", "NNS", "NNP", "NNPS", "VB", "VBD", "VBG", "VBN", "VBP",
             "VBZ"]  # These are the nouns, verbs, and adjectives we want to keep
punctuation = '~><?'';:\/|{}[]@6&*-_,.'
# iterate through the part of speech tags and strip out the conjunction and prepositions.
# keep the verbs, nouns, and adjectives
for i, line in enumerate(resumePOSTags):
    for w in line:
        if w[1] not in POSToKeep:
            resumePOSTags[i].remove(w)

resumeNoStopWordsUpdated = [[w[0] for w in line] for line in resumePOSTags]

# snippet that strips punctuation from the words in the resume.
for line in resumeNoStopWordsUpdated:
    for word in line:
        if word[-1] in punctuation:
            word = word.replace(word[-1], "")

# they layout of this dictionary will be the same as jobsNoStopWords8444
# this has certain parts of speech eliminated
jobsNoStopWordsUpdated = {}

# do the same for the jobs
# something is wrong with this loop (its stripping words we don't want but not all of them)
for i in range(len(jobsNoStopWords)):

    jobsNoStopWordsUpdated[i] = {}
    jobsNoStopWordsUpdated[i]["title"] = jobsNoStopWords[i]["title"]

    # store the job part of speech
    jobPOS = nltk.pos_tag(jobsNoStopWords[i]["description"])
    for w in jobPOS:
        if w[1] not in POSToKeep:
            print("found %s" % w[1])
            jobPOS.remove(w)
    jobsNoStopWordsUpdated[i]["description"] = [w[0] for w in jobPOS]

# Remove special punctuation from the last indice in the jobsNoStopWordsUpdated dictionary description key
for i in range(len(jobsNoStopWordsUpdated)):
    for j in range(len(jobsNoStopWordsUpdated[i]["description"])):
        if jobsNoStopWordsUpdated[i]["description"][j][-1] in punctuation:
            jobsNoStopWordsUpdated[i]["description"][j] = jobsNoStopWordsUpdated[i]["description"][j].replace(
                jobsNoStopWordsUpdated[i]["description"][j][-1], "")

# Done with preprocessing
# the query will be the job description
# now we will run tf-idf on both the job descriptions and the resume & strip the stopwords
# for now assume the resume will be the query as i have currently one resume and multiple job descriptions
# consider concatinating the job title to the description....
print("Performing tf-idf for jobs")

# strip the stopwords
tfidf_vectorizer = TfidfVectorizer(use_idf=False, sublinear_tf=False, stop_words=stopwords.words('english'))
corpus = []
for i in range(len(jobsNoStopWordsUpdated)):
    # concatinate the job title to the description
    corpus.append(
        jobsNoStopWordsUpdated[i]["title"] + " " + " ".join(word for word in jobsNoStopWordsUpdated[i]["description"]))

# document term matrix (and learn the vocabulary) for the jobs
x = tfidf_vectorizer.fit_transform(corpus)
features = tfidf_vectorizer.get_feature_names()  # get the vocabulary from the job dataset

print("tf-idf complete! \n")

# Do the same for the resume...
print("performing tf-idf for the resume (query)")
# resumeCorpus = " ".join(word for word in resumeNoStopWordsUpdated)
resumeCorpus = [" ".join(word for word in section) for section in resumeNoStopWordsUpdated]
# document term matrix for resume
y = tfidf_vectorizer.transform(resumeCorpus)

# run the cosine similarity
similarity = cosine_similarity(x,
                               y).flatten()  # with paragraphs it returns 10 paragraphs * 34 jobs to make a product of 340 scores

# sort to find the related documents based on highest cosine similarity score
# this is for the top 5 scores
# top_5_related_document_indices = similarity.argsort()[:-5:-1]
#
# # This is for all similarity scores from highest to lowest
# all_related_document_indices = similarity.argsort()[::-1]
# top 5 documents with cosine similarity scores are here
# top_5 = {}
# for i in top_5_related_document_indices:
#     top_5[i] = (similarity[i], corpus[i])
#
# id 2 word dictionary taken from the tf idf vectorizer
id2word = {}
for word, id in tfidf_vectorizer.vocabulary_.items():
    id2word[id] = word
